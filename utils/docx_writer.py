from docx import Document
import re


def write_formatted_docx(text: str, output_path: str) -> None:
    """Create a DOCX file with simple formatting for headings and bullet lists."""
    doc = Document()

    bullet_regex = re.compile(r"^(?:[-*\u2022]|\d+[.)])\s+")

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if bullet_regex.match(stripped):
            content = bullet_regex.sub("", stripped)
            doc.add_paragraph(content, style="List Bullet")
        elif stripped.isupper() and len(stripped.split()) <= 6:
            doc.add_heading(stripped, level=2)
        else:
            doc.add_paragraph(stripped)

    doc.save(output_path)
